import cv2
import docx
import smtplib, ssl
from docx.shared import Inches
from email import encoders
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText

def paste_picture(login, password):
    path_input_picture="./pic/change.png"
    path_output_picture = './pic/win10pppoe_7.png'
    red = (0, 0, 255)
    login_position = (212, 144)
    password_position = (212, 181)

    img = cv2.imread(path_input_picture)
    font = cv2.FONT_HERSHEY_COMPLEX
    cv2.putText(img, login, login_position, font, 0.7, color=red, thickness=1)
    cv2.putText(img, password, password_position, font, 0.7, color=red, thickness=1)
    cv2.imwrite(path_output_picture, img)
    cv2.waitKey()

def make_doc(login, password):
    doc = docx.Document()
    images_path = ["./pic/win10pppoe_{}.png".format(i) for i in range(1,10)]
    texts = ['Нажимаем Пуск, затем переходим в раздел "параметры":',
             'Следом выбираем "Сеть и Интернет":',
             'В левом столбце нужно выбрать "Набор номера":',
             'Выбираем Настройка нового подключения:',
             'Выбираем пункт «Подключение к Интернету», нажимаем кнопку «Далее»:',
             'Выбираем «Высокоскоростное (с PPPoE)»',
             'Вводим «Имя пользователя» и «Пароль» (можно скопировать под картинкой). Ставим галочку в поле «Запомнить этот пароль». Затем нажимаем «Подключить»',
             'Ожидаем пока пройдет проверка.',
             "Настройка завершена! Можете пользоваться!"]

    sections = doc.sections
    for section in sections:
        section.top_margin  = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.4)
        section.right_margin = Inches(0.4)

    paste_picture(login, password)
    doc.add_heading('Настройка подключения Win10 (PPPOE) .',0)
    doc.add_paragraph('Настройка подключения к сети Интернет на ПК пользователя под управлением Windows 10' )
    doc.add_paragraph('')

    for i in range(0,9):
        doc.add_paragraph(texts[i])
        doc.add_picture(images_path[i], width=Inches(6))
        if i!=6:
            doc.add_paragraph('')
        if i%2!=0 and i!=6:
            doc.add_page_break()
        if i==6:
            doc.add_paragraph('Ваши учетные данные:')
            doc.add_paragraph('Логин: {}'.format(login))
            doc.add_paragraph('Пароль: {}'.format(password))
    doc.save('manual.docx')

def mail_sender(sender_name, password, sender_email, recipient_email, recipient_name, provider_name):

    text_template = """\nДобрый день, %friend_name%! Это %my_name% от интернет провайдера ""%provider_name%"".
            В продолжении нашего разговора высылаю вам инструкцию по настройке интернета, файл во вложении, там же ваши учетные данные:"""
    text_of_mail = text_template.replace("%friend_name%", recipient_name).replace("%my_name%", sender_name).replace("%provider_name%", provider_name)

    theme_of_mail = "Инструкция по по настройке интернета"

    message = MIMEMultipart()
    message["From"] = sender_email
    message["To"] = recipient_email
    message["Subject"] = theme_of_mail
    message.attach(MIMEText(text_of_mail, "plain"))

    filename = 'manual.docx'
    with open(filename, "rb") as attachment:
        part = MIMEBase("application", "octet-stream")
        part.set_payload(attachment.read())
    encoders.encode_base64(part)
    part.add_header("Content-Disposition", f"attachment; filename= {filename}",)
    message.attach(part)
    text = message.as_string()
    context = ssl.create_default_context()

    with smtplib.SMTP_SSL("smtp.yandex.ru", 465, context=context) as server:
        server.login(sender_email, password)
        server.sendmail(sender_email, recipient_email, text)




